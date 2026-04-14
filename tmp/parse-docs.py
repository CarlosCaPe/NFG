import docx, os

files = [
    'clients/oncohealth/output/downloads/NewUM_Change-Request.docx',
    'clients/oncohealth/output/downloads/NewUM-Data-System-Design.docx'
]

for f in files:
    print()
    print('=' * 80)
    print(f'FILE: {f} ({os.path.getsize(f)} bytes)')
    print('=' * 80)
    doc = docx.Document(f)
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip():
            style = p.style.name if p.style else 'None'
            print(f'  [{style}] {p.text.strip()[:300]}')
    for ti, table in enumerate(doc.tables):
        print()
        print(f'  --- TABLE {ti+1} ({len(table.rows)} rows x {len(table.columns)} cols) ---')
        for ri, row in enumerate(table.rows):
            cells = [c.text.strip()[:60] for c in row.cells]
            sep = ' | '
            print(f'    R{ri}: {sep.join(cells)}')
            if ri > 80:
                remaining = len(table.rows) - ri - 1
                print(f'    ... ({remaining} more rows)')
                break
